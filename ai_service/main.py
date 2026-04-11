import io
import re
import tempfile
import docx
import pdfplumber
from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from thefuzz import process, fuzz
from sentence_transformers import SentenceTransformer, util
from transformers import pipeline
from sqlalchemy.orm import Session
from database import get_db
from models import Candidate, CandidateSkill, Skill

app = FastAPI(title="NexHire AI Service")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load the Sentence Transformer model
model = SentenceTransformer('all-MiniLM-L6-v2')

# Load the QA pipeline for NLP extraction
try:
    qa_pipeline = pipeline("question-answering", model="deepset/roberta-base-squad2")
except Exception as e:
    qa_pipeline = None

# Define normalized skills dictionary
SKILL_MAPPING = {
    "react.js": "React",
    "js": "JavaScript",
    "node.js": "Node.js",
    "ts": "TypeScript",
    "py": "Python",
    "golang": "Go",
    "aws": "AWS",
    "docker": "Docker",
    "k8s": "Kubernetes"
}

STANDARD_SKILLS = [
    "JavaScript", "Python", "React", "Node.js", "Java", "C++",
    "AWS", "Docker", "Kubernetes", "SQL", "MongoDB", "TypeScript",
    "Go", "HTML", "CSS", "Machine Learning", "Data Analysis"
]

# --- Agent 1: Resume Parsing ---

def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text

import zipfile
import xml.etree.ElementTree as ET

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Robust extraction reading all text directly from DOCX XML to overcome nested elements and shapes"""
    try:
        text = []
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as docx_zip:
            # We specifically read the main document xml
            xml_content = docx_zip.read('word/document.xml')
            tree = ET.fromstring(xml_content)
            
            # The namespace for paragraph text is typically w:t
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Find all <w:p> elements (paragraphs)
            for p in tree.findall('.//w:p', namespaces):
                para_text = []
                # Find all <w:t> elements (text runs) within this paragraph
                for t in p.findall('.//w:t', namespaces):
                    if t.text:
                        para_text.append(t.text)
                
                if para_text:
                    text.append("".join(para_text))
                    
        return "\n".join(text)
    except Exception as e:
        # Fallback to python-docx if zipfile parsing fails for some reason
        doc = docx.Document(io.BytesIO(file_bytes))
        fullText = []
        for para in doc.paragraphs:
            if para.text.strip():
                fullText.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            fullText.append(para.text.strip())
        return "\n".join(fullText)

def parse_resume_text(text: str, filename: str = ""):
    """Extract key details using Natural Language Processing (QA model + heuristics)"""
    # Extract Email
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    emails = re.findall(email_pattern, text)
    email = emails[0] if emails else ""
    
    name = "Unknown Candidate"
    experience = 0
    found_skills = []
    
    # Text truncation for QA (LLM context window)
    context_text = text[:3000] # First 3000 characters typically hold name and summary
    
    if qa_pipeline:
        try:
            # Predict Name
            ans_name = qa_pipeline(question="What is the name of the candidate?", context=context_text)
            if ans_name['score'] > 0.3:
                name = ans_name['answer'].title()
            
            # Predict Experience
            ans_exp = qa_pipeline(question="How many years of experience does the candidate have? answer in a number", context=context_text[:2000])
            if ans_exp['score'] > 0.1:
                # regex to find the number in the answer
                exp_match = re.search(r'(\d+)', ans_exp['answer'])
                if exp_match:
                    experience = int(exp_match.group(1))
        except Exception:
            pass

    # Fallback and Heuristics for Name
    if name == "Unknown Candidate" or len(name) < 3 or "@" in name:
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        ignore_keywords = {"resume", "curriculum vitae", "cv", "page", "profile", "contact", "email", "phone", "summary"}
        for line in lines[:20]:
            clean_line = line.strip()
            lower_line = clean_line.lower()
            if len(clean_line) < 3 or len(clean_line.split()) > 4: continue
            if "@" in clean_line or "www." in lower_line or ".com" in lower_line or re.search(r'\d', clean_line): continue
            if any(kw == w for kw in ignore_keywords for w in re.split(r'\W+', lower_line)): continue
            if re.match(r'^[a-zA-Z\s\.\'-]+$', clean_line) and len(clean_line.split()) > 1:
                name = clean_line.title()
                break

    if name == "Unknown Candidate":
        if filename:
            import os
            base_name = os.path.splitext(filename)[0]
            name = base_name.replace('_', ' ').replace('-', ' ').title()
        elif email:
            name = email.split('@')[0].replace('.', ' ').title()

    # NLP and Heuristic Skills Extraction
    text_lower = text.lower()
    for skill in STANDARD_SKILLS:
        # Check whole word match for skills
        if re.search(r'\b' + re.escape(skill.lower()) + r'\b', text_lower):
            found_skills.append(skill)

    # Fallback for experience if QA failed
    if experience == 0:
        experience_pattern = r'(\d+)\+?\s*(?:years|yrs)\s*(?:of\s*)?(?:experience|exp)?'
        exp_matches = re.search(experience_pattern, text_lower)
        if exp_matches:
            experience = int(exp_matches.group(1))

    return {
        "name": name,
        "email": email,
        "skills": found_skills,
        "experience": experience,
        "raw_text": text
    }


# --- Agent 2: Skill Normalization ---

def normalize_skills(skills: list) -> list:
    normalized = []
    for skill in skills:
        skill_lower = skill.lower()
        if skill_lower in SKILL_MAPPING:
            normalized.append(SKILL_MAPPING[skill_lower])
        else:
            # Fuzzy match against standard skills
            match = process.extractOne(skill, STANDARD_SKILLS, scorer=fuzz.token_sort_ratio)
            if match and match[1] > 80:  # 80 is the fuzzy score threshold
                normalized.append(match[0])
            else:
                normalized.append(skill)
    return list(set(normalized))


# --- Agent 3: Matching ---

class MatchRequest(BaseModel):
    job_description: str

def compute_match_score(candidate_text: str, job_description: str) -> float:
    # 1. Semantic similarity using a snippet of the candidate text
    # Limiting to 1000 chars ensures we don't dilute the embedding too much
    embeddings1 = model.encode(candidate_text[:1000], convert_to_tensor=True)
    embeddings2 = model.encode(job_description, convert_to_tensor=True)
    
    cosine_score = util.cos_sim(embeddings1, embeddings2).item()
    
    # 2. Keyword matching boost
    job_lower = job_description.lower()
    cand_lower = candidate_text.lower()
    
    boost = 0.0
    matched_skills = 0
    
    for skill in STANDARD_SKILLS:
        if skill.lower() in job_lower:
            # If the job asks for a standard skill, check if candidate has it
            if skill.lower() in cand_lower:
                boost += 25.0  # Big boost for explicit skill match
                matched_skills += 1
            else:
                boost -= 10.0  # Penalty for missing a requested skill
                
    # If job description mentions no standard skills, just rely more on semantic
    base_score = max(0, cosine_score * 100)
    
    # We want a base score to be roughly 50-60 if it's somewhat relevant
    if base_score > 15:
        base_score += 40
        
    final_score = base_score + boost
    
    # Ensure realistic capping
    return max(0.0, min(final_score, 100.0))


# --- API Routes ---

@app.post("/parse")
async def parse_resume(file: UploadFile = File(...)):
    if not file.filename.endswith(('.pdf', '.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")
    
    try:
        contents = await file.read()
        
        if file.filename.endswith('.pdf'):
            text = extract_text_from_pdf(contents)
        else:
            text = extract_text_from_docx(contents)
            
        parsed_data = parse_resume_text(text, file.filename)
        
        # apply normalization agent
        normalized_skills = normalize_skills(parsed_data['skills'])
        parsed_data['skills'] = normalized_skills
        
        # Return parsed data, let the Node.js backend handle database persistence
        return parsed_data
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

class MathRequestWithCandidates(BaseModel):
    jobDescription: str
    candidates: list

@app.post("/match")
async def match_candidates(req: MathRequestWithCandidates):
    try:
        results = []
        for cand in req.candidates:
            cand_text = cand.get('text', '')
            # encode using candidate text directly
            score = compute_match_score(cand_text, req.jobDescription)
            results.append({
                "id": str(cand.get('id')),
                "score": round(score, 2)
            })
        
        return results
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/match-all")
async def match_all_candidates(req: MatchRequest, db: Session = Depends(get_db)):
    try:
        candidates = db.query(Candidate).all()
        results = []
        for cand in candidates:
            # fetch skills
            cand_text = " ".join([s.skill.name for s in cand.skills])
            score = compute_match_score(cand_text, req.job_description)
            results.append({
                "id": str(cand.id),
                "name": cand.name,
                "role": "Candidate", # Placeholder
                "score": round(score, 2),
                "skills": [s.skill.name for s in cand.skills]
            })
        
        results.sort(key=lambda x: x['score'], reverse=True)
        return {"matches": results}
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/candidates")
async def get_candidates(db: Session = Depends(get_db)):
    try:
        candidates = db.query(Candidate).order_by(Candidate.created_at.desc()).all()
        results = []
        for cand in candidates:
            results.append({
                "id": str(cand.id),
                "name": cand.name,
                "email": cand.email,
                "skills": [s.skill.name for s in cand.skills],
                "created_at": cand.created_at
            })
        return {"candidates": results}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)